# app.py
# PDF Editor Ultimate PRO — final version requested
# Features:
# - Input folder / Output folder (saved in AppData per user)
# - Upload optional
# - Page selection (select, deselect, quick ranges, reorder fallback)
# - Delete pages by not selecting them (save new PDF with only selected pages)
# - OCR (Tesseract ell+eng default)
# - OCR preview before DOCX export
# - OCR progress bar (per-page %)
# - One DOCX per PDF (Simple mode) with page dividers: --- PAGE N ---
# - Save DOCX in output folder with same base filename (overwrite if exists)
# - Settings (tesseract path, input/output folders, ocr_lang) saved to %LOCALAPPDATA%/PDFEditorUltimate/config.json

import streamlit as st
from PyPDF2 import PdfReader, PdfWriter
from pdf2image import convert_from_bytes
from PIL import Image
import io, re, base64, zipfile, hashlib, os, json, datetime, math, time
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# OCR + docx
import pytesseract
from docx import Document
from docx.shared import Pt

# optional sortable UI
try:
    from streamlit_sortables import sort_items
    SORTABLES_AVAILABLE = True
except Exception:
    SORTABLES_AVAILABLE = False

# ---------------------------
# AppData config
# ---------------------------
APP_NAME = "PDFEditorUltimate"
LOCALAPPDATA = os.getenv("LOCALAPPDATA") or os.path.join(os.path.expanduser("~"), "AppData", "Local")
APPDATA_DIR = os.path.join(LOCALAPPDATA, APP_NAME)
CONFIG_PATH = os.path.join(APPDATA_DIR, "config.json")

DEFAULT_CONFIG = {
    "input_folder": "",
    "output_folder": "",
    "tesseract_path": r"C:\Users\{user}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe".format(user=os.getenv("USERNAME") or ""),
    "ocr_lang": "ell+eng"
}

os.makedirs(APPDATA_DIR, exist_ok=True)

def load_config():
    if os.path.exists(CONFIG_PATH):
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                cfg = json.load(f)
            # ensure keys
            for k, v in DEFAULT_CONFIG.items():
                if k not in cfg:
                    cfg[k] = v
            return cfg
        except Exception:
            return DEFAULT_CONFIG.copy()
    else:
        cfg = DEFAULT_CONFIG.copy()
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(cfg, f, indent=2, ensure_ascii=False)
        return cfg

def save_config(cfg):
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(cfg, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        st.error(f"Αποτυχία αποθήκευσης ρυθμίσεων: {e}")
        return False

config = load_config()

# ---------------------------
# Streamlit page config
# ---------------------------
st.set_page_config(page_title="PDF Editor • Ultimate PRO", layout="wide")
st.markdown("""
<style>
:root{ --accent:#4a90e2; --muted:#6b7280; --card:#ffffff; --bg:#f6f8fa;}
.stApp { background: var(--bg); }
.header { background: linear-gradient(90deg, #4a90e2, #06a77d); padding:12px; border-radius:8px; color:white; }
.card { background:var(--card); padding:12px; border-radius:8px; box-shadow:0 1px 6px rgba(16,24,40,0.06); margin-bottom:10px;}
.small { font-size:0.85rem; color:var(--muted); }
textarea { width:100%;}
</style>
""", unsafe_allow_html=True)

st.markdown("<div class='header'><h2>🌈 PDF Editor • Ultimate PRO — OCR → DOCX (ell+eng)</h2></div>", unsafe_allow_html=True)
st.write("Τοπική επεξεργασία — Input/Output folders, επιλογή σελίδων, OCR preview, progress bar, και ένα .docx ανά PDF.")

# ---------------------------
# Sidebar for settings
# ---------------------------
st.sidebar.header("Ρυθμίσεις εφαρμογής (αποθηκεύονται σε AppData)")

st.sidebar.markdown("**Input folder** — όπου βάζεις τα PDFs (προαιρετικό)")
input_folder = st.sidebar.text_input("Input folder (ή leave blank)", value=config.get("input_folder", ""))
if st.sidebar.button("Browse input folder"):
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk(); root.withdraw()
        sel = filedialog.askdirectory(initialdir=input_folder or os.path.expanduser("~"))
        if sel:
            input_folder = sel
    except Exception:
        st.sidebar.info("Δεν υποστηρίζεται ο διάλογος εδώ — γράψε το path χειροκίνητα.")

st.sidebar.markdown(f"**Current:** {input_folder or '— none —'}")

st.sidebar.markdown("**Output folder** — εκεί θα αποθηκεύονται τελικά PDF & DOCX (no timestamp)")
output_folder = st.sidebar.text_input("Output folder (or leave blank)", value=config.get("output_folder",""))
if st.sidebar.button("Browse output folder"):
    try:
        import tkinter as tk
        from tkinter import filedialog
        root = tk.Tk(); root.withdraw()
        sel = filedialog.askdirectory(initialdir=output_folder or os.path.expanduser("~"))
        if sel:
            output_folder = sel
    except Exception:
        st.sidebar.info("Δεν υποστηρίζεται ο διάλογος εδώ — γράψε το path χειροκίνητα.")

st.sidebar.markdown(f"**Current:** {output_folder or '— none —'}")

st.sidebar.markdown("**Tesseract exe (leave blank if in PATH)**")
tess_path_input = st.sidebar.text_input("Path to tesseract.exe", value=config.get("tesseract_path",""))
st.sidebar.markdown("**OCR language (Tesseract)**")
ocr_lang = st.sidebar.text_input("OCR languages, e.g. 'ell+eng'", value=config.get("ocr_lang","ell+eng"))

if st.sidebar.button("Save settings to AppData"):
    cfg_new = config.copy()
    cfg_new["input_folder"] = input_folder.strip()
    cfg_new["output_folder"] = output_folder.strip()
    cfg_new["tesseract_path"] = tess_path_input.strip()
    cfg_new["ocr_lang"] = ocr_lang.strip() or "ell+eng"
    if save_config(cfg_new):
        config.update(cfg_new)
        st.sidebar.success("Ρυθμίσεις αποθηκεύτηκαν.")

# apply tesseract path from config or sidebar
tesseract_path_effective = tess_path_input.strip() or config.get("tesseract_path") or ""
if tesseract_path_effective:
    try:
        pytesseract.pytesseract.tesseract_cmd = tesseract_path_effective
    except Exception:
        pass

# ---------------------------
# Helpers
# ---------------------------
def sanitize_filename(name: str):
    return re.sub(r'[^A-Za-z0-9_\-\.Α-Ωα-ωίόάέύήϊϋΰΪΫ ]', '_', name)

def convert_pdf_to_images(pdf_bytes: bytes, dpi=200):
    try:
        images = convert_from_bytes(pdf_bytes, dpi=dpi)
        return images
    except Exception as e:
        st.error(f"convert_from_bytes failed: {e}")
        return []

def ocr_image_to_text(img: Image.Image, lang="ell+eng"):
    try:
        return pytesseract.image_to_string(img, lang=lang)
    except Exception as e:
        return ""

def extract_embedded_text(pdf_bytes: bytes):
    text = ""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        for p in reader.pages:
            try:
                t = p.extract_text() or ""
            except Exception:
                t = ""
            text += "\n--- PAGE ---\n" + t + "\n"
    except Exception:
        return ""
    return text

def write_selected_pages_pdf(pdf_bytes: bytes, selected_indices, out_path):
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        writer = PdfWriter()
        for i in selected_indices:
            if i < len(reader.pages):
                writer.add_page(reader.pages[i])
        # write to disk overwriting if exists
        with open(out_path, "wb") as f:
            writer.write(f)
        return True
    except Exception as e:
        st.error(f"Saving PDF failed: {e}")
        return False

def text_to_docx_simple(text: str, title: str = None):
    """Simple docx export (one file per PDF) with page dividers --- PAGE N ---"""
    doc = Document()
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(14)
        doc.add_paragraph()
    # write text as-is preserving blank lines
    for line in text.splitlines():
        # keep as paragraph per line
        p = doc.add_paragraph()
        p.add_run(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def search_in_text(text: str, terms):
    results = {}
    lower = text.lower()
    for term in terms:
        if not term: continue
        t = term.lower()
        count = lower.count(t)
        contexts = []
        if count > 0:
            start = 0
            for _ in range(min(count, 5)):
                idx = lower.find(t, start)
                if idx == -1:
                    break
                s = max(0, idx-30)
                e = min(len(text), idx + len(t) + 30)
                contexts.append(text[s:e].replace("\n"," "))
                start = idx + len(t)
        results[term] = {"count": count, "contexts": contexts}
    return results

# ---------------------------
# File collection
# ---------------------------
st.markdown("### Επιλογή / Ανέβασμα PDF")
st.write("Μπορείς να ανεβάσεις ένα ή πολλά PDF ή να τα τοποθετήσεις στον φάκελο εισόδου που όρισες.")

uploaded = st.file_uploader("Ανέβασε ένα ή πολλά PDF (ή άφησέ τα στον input folder)", accept_multiple_files=True, type="pdf")
files_to_process = []

# If input folder provided in session or config, prefer it when no upload
session_input = st.session_state.get("input_folder") if "input_folder" in st.session_state else config.get("input_folder","")
effective_input = uploaded and uploaded or []
# prefer uploaded files if present
if uploaded:
    files_to_process = uploaded
else:
    input_folder_effective = input_folder.strip() or config.get("input_folder","")
    if input_folder_effective:
        try:
            pdf_list = [f for f in os.listdir(input_folder_effective) if f.lower().endswith(".pdf")]
            if pdf_list:
                for name in pdf_list:
                    path = os.path.join(input_folder_effective, name)
                    with open(path, "rb") as fh:
                        data = fh.read()
                    # create simple object with .name and .read
                    class SimpleFile:
                        def __init__(self, name, data):
                            self.name = name
                            self._data = data
                        def read(self):
                            return self._data
                        def seek(self, pos):
                            pass
                    files_to_process.append(SimpleFile(name, data))
            else:
                st.info("Δεν βρέθηκαν PDF στον input folder.")
        except Exception as e:
            st.error(f"Σφάλμα ανάγνωσης input folder: {e}")

if not files_to_process:
    st.stop()

# Global actions
st.markdown("---")
st.markdown("### Global actions")
col1, col2, col3 = st.columns([1,1,1])
with col1:
    if st.button("Process all (OCR → DOCX)"):
        st.session_state["_process_all"] = True
with col2:
    if st.button("Preview first file (OCR/extract)"):
        st.session_state["_preview_first"] = True
with col3:
    if st.button("Search across all (terms)"):
        st.session_state["_search_all"] = True

search_input = st.text_input("Αναζήτηση (λέξεις/φράσεις, κόμμα)", value="")
search_terms = [s.strip() for s in search_input.split(",") if s.strip()]

# processed outputs list for downloads
processed_outputs = []

# iterate files
for file_obj in files_to_process:
    name = getattr(file_obj, "name", "uploaded.pdf")
    st.markdown("---")
    st.subheader(name)

    # read bytes
    try:
        pdf_bytes = file_obj.read()
    except Exception:
        try:
            pdf_bytes = file_obj.getbuffer().tobytes()
        except Exception as e:
            st.error(f"Cannot read {name}: {e}")
            continue

    # build reader
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        total_pages = len(reader.pages)
    except Exception:
        reader = None
        total_pages = 0

    # thumbnails generation (low-res) - try/except
    images = []
    try:
        images = convert_pdf_to_images(pdf_bytes, dpi=100)
    except Exception:
        images = [None] * total_pages

    # initialize selection state for this file
    key_sel = f"sel_{name}"
    if key_sel not in st.session_state:
        st.session_state[key_sel] = [True] * total_pages if total_pages>0 else []

    # quick selection input and buttons
    sel_col1, sel_col2, sel_col3 = st.columns([1,1,2])
    with sel_col1:
        if st.button(f"Select All [{name}]", key=f"selectall_{name}"):
            st.session_state[key_sel] = [True]*total_pages
    with sel_col2:
        if st.button(f"Deselect All [{name}]", key=f"deselectall_{name}"):
            st.session_state[key_sel] = [False]*total_pages
    with sel_col3:
        quick_input = st.text_input(f"Quick selection (e.g. 2,5,8 or 3-6 or -3 to unselect) for {name}", key=f"quick_{name}")

        if st.button("Apply quick", key=f"applyquick_{name}"):
            txt = quick_input.strip()
            def parse_parts(txt):
                nums = []
                for part in txt.split(","):
                    p = part.strip()
                    if not p: continue
                    neg = False
                    if p.startswith("-"):
                        neg = True; p = p[1:]
                    if "-" in p:
                        a,b = p.split("-",1)
                        a=int(a); b=int(b)
                        for n in range(a,b+1):
                            nums.append((n,neg))
                    else:
                        nums.append((int(p),neg))
                return nums
            try:
                changes = parse_parts(txt)
                for n,neg in changes:
                    if 1 <= n <= total_pages:
                        st.session_state[key_sel][n-1] = (not neg)
                st.success("Applied")
            except Exception:
                st.error("Invalid input for quick selection")

    # show thumbnails and checkboxes
    st.markdown("**Thumbnails & selection**")
    if total_pages == 0:
        st.info("No pages detected or PDF corrupted.")
    else:
        cols = st.columns(3)
        for i in range(total_pages):
            c = cols[i%3]
            with c:
                if i < len(images) and images[i] is not None:
                    buf = io.BytesIO()
                    try:
                        images[i].save(buf, format="PNG")
                        img_b64 = base64.b64encode(buf.getvalue()).decode()
                        st.markdown(f'<img src="data:image/png;base64,{img_b64}" width="220"/>', unsafe_allow_html=True)
                    except Exception:
                        st.text(f"Page {i+1}")
                else:
                    st.text(f"Page {i+1}")
                # extract small snippet if possible
                snippet = ""
                try:
                    snippet = reader.pages[i].extract_text() or ""
                except Exception:
                    snippet = ""
                snippet = snippet.replace("\n", " ")
                snippet = (snippet[:180] + "...") if len(snippet) > 180 else snippet
                st.markdown(f"<div class='small'>{snippet}</div>", unsafe_allow_html=True)
                sel = st.checkbox(f"Page {i+1}", value=st.session_state[key_sel][i], key=f"{name}_cb_{i}")
                st.session_state[key_sel][i] = sel
            if i%3 == 2:
                cols = st.columns(3)

    # show selected pages summary
    selected_indices = [i for i,k in enumerate(st.session_state[key_sel]) if k]
    st.markdown(f"**Selected pages:** {', '.join(str(i+1) for i in selected_indices) if selected_indices else '— none —'}")

    # reorder (if sortable available) - fallback to text input
    reorder_key = f"order_{name}"
    if not selected_indices:
        st.info("Select pages to enable reorder.")
        st.session_state[reorder_key] = []
    else:
        st.markdown("#### Reorder selected pages")
        # prepare miniature images/labels
        items = []
        for idx in selected_indices:
            if idx < len(images) and images[idx] is not None:
                b = io.BytesIO(); images[idx].save(b, format="PNG")
                items.append(f'<div style="text-align:center"><img src="data:image/png;base64,{base64.b64encode(b.getvalue()).decode()}" width="120"><div>Pg {idx+1}</div></div>')
            else:
                items.append(f"Pg {idx+1}")
        if SORTABLES_AVAILABLE:
            try:
                res = sort_items(items, direction="horizontal", key=f"sort_{name}")
                order = [selected_indices[i] for i in res.get("order", list(range(len(items))))]
                st.session_state[reorder_key] = order
            except Exception:
                st.warning("Drag & drop failed; use text input below.")
                st.session_state[reorder_key] = selected_indices
        else:
            st.info("Drag & drop not installed; use text input for order.")
            st.session_state[reorder_key] = selected_indices

        order_txt = st.text_input(f"Alternative order (e.g. {','.join(str(i+1) for i in selected_indices)})", key=f"ordertxt_{name}")
        if order_txt:
            try:
                new = [int(x.strip())-1 for x in order_txt.split(",") if x.strip().isdigit()]
                if set(new) == set(selected_indices):
                    st.session_state[reorder_key] = new
                else:
                    st.warning("New order must include exactly the same selected pages.")
            except Exception:
                st.error("Invalid format for order.")

    # Export area for this file
    st.markdown("### Export / OCR / Save")
    e1, e2, e3, e4 = st.columns([1,1,1,1])
    with e1:
        if st.button(f"Save PDF with selected pages — {name}", key=f"savepdf_{name}"):
            if not selected_indices:
                st.error("No pages selected")
            else:
                order = st.session_state.get(reorder_key, selected_indices)
                # prepare writer with order
                reader_local = PdfReader(io.BytesIO(pdf_bytes))
                writer = PdfWriter()
                for idx in order:
                    writer.add_page(reader_local.pages[idx])
                # write to output folder with same filename (overwrite)
                out_folder = output_folder.strip() or config.get("output_folder","") or ""
                if not out_folder:
                    st.error("Output folder not set in sidebar or saved config.")
                else:
                    os.makedirs(out_folder, exist_ok=True)
                    out_path = os.path.join(out_folder, name)
                    try:
                        with open(out_path, "wb") as f:
                            writer.write(f)
                        st.success(f"Saved PDF to {out_path}")
                    except Exception as e:
                        st.error(f"Save failed: {e}")

    with e2:
        if st.button(f"OCR Preview (selected pages) — {name}", key=f"ocr_preview_{name}"):
            if not selected_indices:
                st.error("No pages selected for OCR preview")
            else:
                # convert only selected pages to images (to save time)
                with st.spinner("Converting selected pages to images..."):
                    try:
                        imgs = convert_pdf_to_images(pdf_bytes, dpi=300)
                    except Exception as e:
                        st.error(f"Image conversion failed: {e}")
                        imgs = []
                if not imgs:
                    st.error("Cannot obtain page images for OCR.")
                else:
                    # build preview text only for selected pages and show first 1000 chars
                    preview_text = ""
                    for i_idx, page_idx in enumerate(selected_indices):
                        if page_idx < len(imgs):
                            txt = ocr_image_to_text(imgs[page_idx], lang=ocr_lang)
                        else:
                            txt = ""
                        preview_text += f"\n--- PAGE {page_idx+1} ---\n{txt}\n"
                    st.markdown("**OCR Preview (first 1000 chars):**")
                    st.code((preview_text[:1000] + ("..." if len(preview_text) > 1000 else "")))
                    # store preview in session in case user wants to export docx immediately
                    st.session_state[f"preview_text_{name}"] = preview_text

    with e3:
        if st.button(f"OCR -> DOCX (selected pages) — {name}", key=f"ocr_docx_{name}"):
            # perform OCR on selected pages, with progress bar, create one docx per PDF
            if not selected_indices:
                st.error("No pages selected")
            else:
                # set tesseract path from settings if provided
                if tesseract_path_effective:
                    try:
                        pytesseract.pytesseract.tesseract_cmd = tesseract_path_effective
                    except Exception:
                        pass
                # convert all pages or only selected pages to images
                with st.spinner("Converting PDF pages to images..."):
                    try:
                        imgs = convert_pdf_to_images(pdf_bytes, dpi=300)
                    except Exception as e:
                        st.error(f"Image conversion failed: {e}")
                        imgs = []
                if not imgs:
                    st.error("No images for OCR.")
                else:
                    total = len(selected_indices)
                    progress_bar = st.progress(0)
                    full_text = ""
                    for idx_i, page_idx in enumerate(selected_indices):
                        # avoid index error
                        if page_idx < len(imgs):
                            try:
                                txt = ocr_image_to_text(imgs[page_idx], lang=ocr_lang)
                            except Exception:
                                txt = ""
                        else:
                            txt = ""
                        full_text += f"\n--- PAGE {page_idx+1} ---\n{txt}\n"
                        progress = int(((idx_i+1)/total) * 100)
                        progress_bar.progress(progress)
                    progress_bar.empty()
                    # show preview
                    st.markdown("**OCR Result preview (first 1000 chars):**")
                    st.code((full_text[:1000] + ("..." if len(full_text) > 1000 else "")))
                    # save docx (simple mode)
                    out_folder = output_folder.strip() or config.get("output_folder","") or ""
                    if not out_folder:
                        st.error("Output folder not set in sidebar or saved config.")
                    else:
                        os.makedirs(out_folder, exist_ok=True)
                        base = os.path.splitext(name)[0]
                        out_name = f"{base}.docx"   # no timestamp, overwrite allowed
                        out_path = os.path.join(out_folder, out_name)
                        try:
                            buf = text_to_docx_simple(full_text, title=name)
                            with open(out_path, "wb") as f:
                                f.write(buf.getvalue())
                            st.success(f"Saved DOCX to {out_path}")
                            processed_outputs.append((out_name, io.BytesIO(buf.getvalue())))
                        except Exception as e:
                            st.error(f"Save DOCX failed: {e}")

    with e4:
        if st.button(f"Extract embedded text -> DOCX (selected pages) — {name}", key=f"emb_docx_{name}"):
            # try to extract embedded text from pages and save docx
            if not selected_indices:
                st.error("No pages selected")
            else:
                # build text only for selected pages
                full_text = ""
                try:
                    reader_local = PdfReader(io.BytesIO(pdf_bytes))
                    for pidx in selected_indices:
                        try:
                            t = reader_local.pages[pidx].extract_text() or ""
                        except Exception:
                            t = ""
                        full_text += f"\n--- PAGE {pidx+1} ---\n{t}\n"
                    # save
                    out_folder = output_folder.strip() or config.get("output_folder","") or ""
                    if not out_folder:
                        st.error("Output folder not set in sidebar or saved config.")
                    else:
                        os.makedirs(out_folder, exist_ok=True)
                        base = os.path.splitext(name)[0]
                        out_name = f"{base}.docx"
                        out_path = os.path.join(out_folder, out_name)
                        buf = text_to_docx_simple(full_text, title=name)
                        with open(out_path, "wb") as f:
                            f.write(buf.getvalue())
                        st.success(f"Saved DOCX to {out_path}")
                        processed_outputs.append((out_name, io.BytesIO(buf.getvalue())))
                except Exception as e:
                    st.error(f"Embedded extraction failed: {e}")

    # Manual download original PDF
    try:
        st.download_button("Download original PDF", data=pdf_bytes, file_name=name, mime="application/pdf")
    except Exception:
        pass

# Final downloads
if processed_outputs:
    st.markdown("---")
    st.header("Downloads")
    for fname, buf in processed_outputs:
        buf.seek(0)
        st.download_button(f"Download: {fname}", data=buf.getvalue(), file_name=fname, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("Δεν δημιουργήθηκαν .docx αρχεία ακόμα.")

st.markdown("---")
st.markdown(f"© PDF Editor Ultimate • Local-only • Config stored in {CONFIG_PATH}")